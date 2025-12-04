import streamlit as st
from pdf2docx import Converter
import os
import tempfile
import time
from docx import Document

# --- 1. Config (Clean & Pro) ---
st.set_page_config(page_title="PDF2Word Pro", page_icon="📄", layout="centered")

st.markdown("""
    <style>
        .block-container { padding-top: 2rem; padding-bottom: 2rem; }
        .stButton>button { 
            width: 100%; 
            background-color: #000000; 
            color: white; 
            font-weight: bold; 
            border-radius: 8px; 
            height: 50px;
        }
        /* Style ให้ UI ดูสะอาดตา */
        .stAlert { padding: 0.5rem; }
        div[data-testid="column"] { gap: 0.5rem; }
    </style>
""", unsafe_allow_html=True)

# --- 2. Logic (Stable & Complete) ---
def repair_thai_docx(docx_path):
    try:
        doc = Document(docx_path)
        # Logic ซ่อมสระ ำ (แบบเบาเครื่อง)
        def fix_sara_am(text):
            if not text or " ำ" not in text: return text
            return text.replace(" ำ", "ำ").replace(" ำ", "ำ")

        for para in doc.paragraphs:
            for run in para.runs:
                run.text = fix_sara_am(run.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = fix_sara_am(run.text)
        doc.save(docx_path)
        return True
    except: return False

def convert_pdf_to_docx(uploaded_file, start_page, end_page, status_box, progress_bar):
    with tempfile.TemporaryDirectory() as temp_dir:
        pdf_path = os.path.join(temp_dir, uploaded_file.name)
        with open(pdf_path, "wb") as f: f.write(uploaded_file.getbuffer())
        
        docx_name = os.path.splitext(uploaded_file.name)[0] + ".docx"
        docx_path = os.path.join(temp_dir, docx_name)
        
        try:
            status_box.info("⚙️ เริ่มต้นกระบวนการ... (Initializing)")
            progress_bar.progress(10)
            
            # โหลดไฟล์เพื่อเตรียมแปลง
            cv = Converter(pdf_path)
            
            # ถ้า end_page เป็น None แปลว่าเอาถึงหน้าสุดท้าย
            if end_page is None:
                end_page = len(cv.pages)
            
            status_box.info(f"📄 กำลังแปลงหน้า {start_page} ถึง {end_page} (เก็บรายละเอียดครบถ้วน)...")
            progress_bar.progress(30)
            
            # --- จุดสำคัญ: สั่งแปลงเฉพาะหน้าที่เลือก ---
            # start ต้องลบ 1 เพราะคอมนับเริ่มจาก 0
            # multi_processing=False (สำคัญมาก ห้ามเอาออก เพื่อกันเครื่องค้าง)
            cv.convert(docx_path, start=start_page-1, end=end_page, multi_processing=False)
            
            cv.close()
            
            progress_bar.progress(80)
            status_box.info("🔧 กำลังซ่อมสระภาษาไทย (Fixing Thai Vowels)...")
            repair_thai_docx(docx_path)
            progress_bar.progress(100)
            
            with open(docx_path, "rb") as f: docx_data = f.read()
            return docx_data, docx_name
            
        except Exception as e:
            st.error(f"เกิดข้อผิดพลาด: {e}")
            return None, None

# --- 3. UI (Smart Range Selector) ---

c1, c2 = st.columns([3, 1])
c1.markdown("### 📄 PDF to Word `Smart`")
c2.markdown("<div style='text-align: right; color: gray; font-size: 0.8em; padding-top: 10px;'>v3.0 Hybrid</div>", unsafe_allow_html=True)

st.divider()

uploaded_file = st.file_uploader("Upload PDF file", type="pdf", label_visibility="collapsed")

if uploaded_file:
    # อ่านจำนวนหน้าเบื้องต้น (ใช้ trick อ่านเร็วๆ)
    try:
        from pypdf import PdfReader
        reader = PdfReader(uploaded_file)
        total_pages = len(reader.pages)
    except:
        # ถ้าอ่านไม่ได้ ให้เดาว่ามีเยอะไว้ก่อน
        total_pages = 100 

    # --- ส่วนควบคุมการเลือกหน้า ---
    st.write(f"พบเอกสารทั้งหมด **{total_pages}** หน้า")
    
    col_opt, col_range = st.columns([1, 2])
    
    with col_opt:
        mode = st.radio("ตัวเลือกการแปลง:", ["ทั้งหมด (All)", "ระบุหน้า (Custom)"])
    
    start_p, end_p = 1, None
    
    with col_range:
        if mode == "ระบุหน้า (Custom)":
            c_s, c_e = st.columns(2)
            with c_s:
                start_p = st.number_input("หน้าเริ่ม", min_value=1, max_value=total_pages, value=1)
            with c_e:
                end_p = st.number_input("ถึงหน้า", min_value=start_p, max_value=total_pages, value=min(start_p+4, total_pages))
            st.caption("💡 แนะนำ: แปลงทีละ 5-10 หน้าจะเร็วมาก")
        else:
            st.info(f"⚠️ แปลงรวดเดียว {total_pages} หน้า อาจใช้เวลา 3-5 นาที")

    st.markdown("---")
    
    run_btn = st.button("🚀 เริ่มแปลงไฟล์ (Start Convert)")
    status_box = st.empty()
    progress_bar = st.empty()

    if run_btn:
        start_time = time.time()
        
        docx_data, docx_name = convert_pdf_to_docx(uploaded_file, start_p, end_p, status_box, progress_bar)
        
        duration = time.time() - start_time
        
        if docx_data:
            status_box.success("✅ เสร็จเรียบร้อย!")
            
            c_info, c_btn = st.columns([1.5, 2])
            with c_info:
                st.caption(f"⏱️ เวลา: {duration:.2f}s | 📦 ขนาด: {len(docx_data)/1024:.1f} KB")
            with c_btn:
                st.download_button(
                    label="📥 ดาวน์โหลด Word (.docx)",
                    data=docx_data,
                    file_name=docx_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
else:
    st.markdown(
        """
        <div style='text-align: center; color: #888; padding: 20px;'>
            <div style='font-size: 3em; margin-bottom: 10px;'>📄</div>
            <div>อัปโหลดไฟล์ PDF เพื่อเริ่มต้น</div>
            <div style='font-size: 0.8em; color: #999;'>(รองรับภาษาไทย + เก็บรูปภาพครบถ้วน)</div>
        </div>
        """, 
        unsafe_allow_html=True
    )
